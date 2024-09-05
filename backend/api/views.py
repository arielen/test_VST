import re

from django.http import FileResponse
from django.shortcuts import get_object_or_404
from rest_framework import status
from rest_framework.views import APIView
from rest_framework.generics import GenericAPIView, ListAPIView
from rest_framework.response import Response
from rest_framework.parsers import FormParser, MultiPartParser
from rest_framework.exceptions import NotFound

from drf_spectacular.utils import extend_schema

from .models import UploadedTextFile, WordOccurrence, Word
from .serializers import (
    UploadedTextFileSerializer, UploadedTextFileListSerializer,
    WordStatisticsSerializer
)

from collections import Counter
from typing import Dict, List

from docx import Document


@extend_schema(
    tags=['Uploaded Text Files'],
)
class UploadedTextFileView(GenericAPIView):
    parser_classes = (MultiPartParser, FormParser, )

    def post(self, request, *args, **kwargs) -> Response:
        file = request.FILES.get('file')

        if not file:
            return Response(
                {'error': 'No file provided'},
                status=status.HTTP_400_BAD_REQUEST
            )

        serializer = UploadedTextFileSerializer(
            data={'file': file, 'file_name': file.name}
        )

        if serializer.is_valid():
            serializer.save()
            self.process_file(serializer.instance, file)
            return Response(serializer.data, status=status.HTTP_201_CREATED)

        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    def process_file(self, text_file: UploadedTextFile, file) -> None:
        file_path = text_file.file.path

        if file.name.endswith('.docx'):
            text = self._extract_text_from_docx(file_path)
        else:
            text = self._extract_text_from_txt(file_path)

        words = re.findall(r'\b\w+\b', text.lower())
        word_counts = Counter(words)

        WordOccurrence.objects.bulk_create(
            [
                WordOccurrence(
                    text_file=text_file,
                    word=Word.objects.get_or_create(text=word)[0],
                    count=count
                )
                for word, count in word_counts.items()
            ]
        )

    def _extract_text_from_txt(self, file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()

    def _extract_text_from_docx(self, file_path: str) -> str:
        doc = Document(file_path)
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])


@extend_schema(
    tags=['Word Statistics'],
)
class WordStatisticsListView(ListAPIView):
    serializer_class = WordStatisticsSerializer

    def get_queryset(self) -> List[Word]:
        """
        Retrieves the queryset for the API endpoint. If a pk is provided, the
        queryset is filtered by the file with that pk. If no pk is provided,
        the queryset is for all files.
        """
        pk = self.kwargs.get('pk')
        self.text_file = None

        if pk:
            try:
                self.text_file = UploadedTextFile.objects.get(pk=pk)
            except UploadedTextFile.DoesNotExist:
                raise NotFound(f"No file found with id {pk}")

            return Word.objects.filter(occurrences__text_file=self.text_file).distinct()

        return Word.objects.all()

    def get_serializer_context(self) -> Dict:
        """
        Generates the context for the serializer. The context includes the
        text_file, which is None if no pk was provided.
        """
        context = super().get_serializer_context()
        context['text_file'] = self.text_file
        return context


@extend_schema(
    tags=['Uploaded Text Files'],
)
class UploadedTextFileList(ListAPIView):
    """
    API endpoint for retrieving a list of uploaded text files or a single
    uploaded text file.

    If a pk is provided, the API endpoint will return a single uploaded text
    file with that pk. Otherwise, the API endpoint will return a list of all
    uploaded text files.
    """
    serializer_class = UploadedTextFileListSerializer

    def get_queryset(self) -> List[UploadedTextFile]:
        """
        Retrieves the queryset for the API endpoint. If a pk is provided, the
        queryset is filtered by the file with that pk. If no pk is provided,
        the queryset is for all files.
        """
        pk = self.kwargs.get('pk')

        if pk:
            return UploadedTextFile.objects.filter(id=pk)

        return UploadedTextFile.objects.all()


class FileDownloadMixin:
    """Mixin to handle common file retrieval logic for text and docx files."""

    def get_file_response(self, pk, disposition_type='attachment'):
        text_file = get_object_or_404(UploadedTextFile, pk=pk)

        if not text_file.file:
            return Response({'error': 'File not found.'}, status=status.HTTP_400_BAD_REQUEST)

        file_handle = text_file.file.open('rb')
        content_type = self._get_content_type(text_file.file.name)

        if disposition_type == 'attachment':
            response = FileResponse(file_handle, content_type=content_type)
            response['Content-Disposition'] = f'attachment; filename="{
                text_file.file_name}"'
            return response
        elif disposition_type == 'inline':
            return self._serve_file_inline(text_file, file_handle, content_type)

    def _get_content_type(self, file_name):
        if file_name.endswith('.docx'):
            return 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        elif file_name.endswith('.txt'):
            return 'text/plain'
        return 'application/octet-stream'

    def _serve_file_inline(self, text_file, file_handle, content_type):
        if text_file.file.name.endswith('.docx'):
            doc = Document(file_handle)
            file_content = "\n".join(
                [paragraph.text for paragraph in doc.paragraphs]
            )
            return Response(file_content, content_type='text/plain')
        else:
            return FileResponse(file_handle, content_type=content_type)


@extend_schema(tags=['Download Files'])
class DownloadFileView(APIView, FileDownloadMixin):
    """
    API endpoint for downloading a single uploaded text file.
    """

    def get(self, request, pk, format=None) -> Response:
        """Download the uploaded text file as an attachment."""
        return self.get_file_response(pk, 'attachment')


@extend_schema(tags=['Download Files'])
class ShowDownloadFileView(APIView, FileDownloadMixin):
    """
    API endpoint for showing a single uploaded text file inline.
    """

    def get(self, request, pk, format=None) -> Response:
        """Display the uploaded text file inline in the browser."""
        return self.get_file_response(pk, 'inline')
